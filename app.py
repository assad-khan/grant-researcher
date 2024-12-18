import streamlit as st
from crewai import Agent, Task, Crew
from langchain.chat_models import ChatOpenAI
import os
from docx import Document
from crewai_tools import (
    SerperDevTool,
    ScrapeWebsiteTool,
)
import io
import sys
import re

os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
os.environ["SERPER_API_KEY"] =os.getenv("SERPER_API_KEY")

# Initialize language models
llm = ChatOpenAI(model_name="gpt-4o-mini")

# Define tools
tools = [SerperDevTool(), ScrapeWebsiteTool()]

class StreamToExpander:
    def __init__(self, expander, buffer_limit=10000):
        self.expander = expander
        self.buffer = []
        self.buffer_limit = buffer_limit

    def write(self, data):
        # Clean ANSI escape codes from output
        cleaned_data = re.sub(r'\x1B\[\d+;?\d*m', '', data)
        if len(self.buffer) >= self.buffer_limit:
            self.buffer.pop(0)
        self.buffer.append(cleaned_data)

        if "\n" in data:
            self.expander.markdown(''.join(self.buffer), unsafe_allow_html=True)
            self.buffer.clear()

    def flush(self):
        if self.buffer:
            self.expander.markdown(''.join(self.buffer), unsafe_allow_html=True)
            self.buffer.clear()


def create_agents(llm):
    """Creates and returns agents with predefined roles and goals."""
    try:
        researcher = Agent(
            role='Grant Researcher',
            goal='Find suitable grants for the organization',
            backstory="You are an expert in finding and analyzing grant opportunities.",
            tools=tools,
            verbose=True,
            llm=llm
        )

        analyzer = Agent(
            role='Grant Analyzer',
            goal='Analyze grant requirements and organizational fit',
            backstory="You are an expert in analyzing grant requirements and assessing organizational eligibility.",
            tools=tools,
            verbose=True,
            llm=llm
        )

        writer = Agent(
            role='Grant Writer',
            goal='Write compelling grant applications',
            backstory="You are a skilled grant writer with a track record of successful applications.",
            tools=[],
            verbose=True,
            llm=llm
        )
        return researcher, analyzer, writer
    except Exception as e:
        st.error(f"Error creating agents: {str(e)}")
        return None, None, None

def create_tasks(researcher, analyzer, writer):
    """Creates and returns tasks for each agent."""
    try:
        research_task = Task(
            description="Research and identify suitable grants based on the organization's profile and needs. Organization: {organization}, Mission: {mission}, Project: {project}, Funding Needed: ${funding}. Use the provided website links: {websites} if available, otherwise search randomly.",
            agent=researcher,
            expected_output="A list of at least 3 potential grants with their names, funding amounts, brief descriptions, website links, and application deadlines. Also give a list of Grant Websites Direct URL to grant opportunity"
        )

        analysis_task = Task(
            description="Analyze the identified grants for eligibility and fit with the organization. Also, find and analyze similar successful grant applications from the web. Organization: {organization}, Mission: {mission}, Project: {project}, Funding Needed: ${funding}",
            agent=analyzer,
            expected_output="A detailed analysis of each grant, including eligibility criteria, alignment with organization goals, probability of success, and insights from similar successful applications. Also Grant Application Reference Websites URLs of successful grant application examples"
        )

        writing_task = Task(
            description="Write a compelling grant application for the selected grant opportunity based on the requirements and analysis. Organization: {organization}, Mission: {mission}, Project: {project}, Funding Needed: ${funding}",
            agent=writer,
            expected_output="A comprehensive, detailed grant application draft, including an executive summary, project description, budget overview, expected outcomes, and any specific sections required by the grant guidelines. Also Source Links (URLs) of the target grant"
        )
        return [research_task, analysis_task, writing_task]
    except Exception as e:
        st.error(f"Error creating tasks: {str(e)}")
        return []

def run_grant_process(grant_crew, input_data):
    """Runs the grant research and writing process."""
    try:
        with st.spinner("Processing... This may take a few minutes."):
            result = grant_crew.kickoff(inputs=input_data)
        st.success("Process completed!")
        return result
    except Exception as e:
        st.error(f"An error occurred during the process: {str(e)}")
        return None

def generate_download_link(result):
    """Generates a download link for the result in a Word document format."""
    try:
        from docx import Document  # Ensure docx is imported within the function

        # Create a Word document
        doc = Document()
        doc.add_heading("Grant Research and Writing Results", level=1)

        # Check if `result` is structured as a dictionary or a list
        if isinstance(result, dict):
            for section, content in result.items():
                doc.add_heading(section, level=2)
                doc.add_paragraph(str(content))  # Ensure content is a string
        elif isinstance(result, (list, tuple)):
            for item in result:
                doc.add_paragraph(str(item))  # Convert each item to a string if needed
        else:
            doc.add_paragraph(str(result))  # Ensure `result` is a string if it’s a single output

        # Save document to an in-memory file
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Provide download link for Word document
        st.download_button(
            label="Download Results as Word Document",
            data=buffer,
            file_name="grant_results.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"Error generating download link: {str(e)}")

def main():
    # Streamlit app
    st.title("Automated Grant Research and Writing Assistant")

    # User input
    org_name = st.text_input("Organization Name")
    org_mission = st.text_area("Organization Mission")
    project_description = st.text_area("Project Description")
    funding_amount = st.number_input("Desired Funding Amount", min_value=0)
    websites = st.text_area("Enter specific website links for grant search (optional, one per line)")

    if st.button("Start Grant Research and Writing Process"):
        process_output_expander = st.expander("Processing Output:")
        sys.stdout = StreamToExpander(process_output_expander)
        if org_name and org_mission and project_description and funding_amount:
            input_data = {
                "organization": org_name,
                "mission": org_mission,
                "project": project_description,
                "funding": funding_amount,
                "websites": websites.split('\n') if websites else []
            }

            # Create agents and tasks
            researcher, analyzer, writer = create_agents(llm)
            if researcher and analyzer and writer:
                tasks = create_tasks(researcher, analyzer, writer)
                grant_crew = Crew(agents=[researcher, analyzer, writer], tasks=tasks, verbose=True)

                # Run process and provide download link
                result = run_grant_process(grant_crew, input_data)
                if result:
                    st.subheader("Grant Research and Writing Results")
                    st.markdown(result)

                    # Provide file download option
                    generate_download_link(result)

        else:
            st.error("Please fill in all the required fields before starting the process.")
            
            
if __name__ == "__main__":
    main()
